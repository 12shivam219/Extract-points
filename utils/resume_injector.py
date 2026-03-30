from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
import io
import re
import logging
from .bookmark_manager import BookmarkManager

# Setup logging
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

class ResumeInjector:
    """Handles injecting extracted points into resume templates with bookmarks."""
    
    def __init__(self):
        self.bookmark_manager = BookmarkManager()
    
    def extract_points_by_heading(self, processed_text):
        """
        Extract points from processed text organized by cycle.
        Works with Cycle format:
        Cycle 1: / • Point 1
        Cycle 2: / • Point 2
        etc.
        
        Returns: points_by_cycle dict like {1: [points], 2: [points], ...}
        """
        points_by_cycle = {}
        current_cycle = None
        all_points = []  # Track all points found
        
        lines = processed_text.split('\n')
        
        logger.debug(f"Total lines: {len(lines)}")
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for cycle header (Cycle 1:, Cycle 2:, etc.)
            cycle_match = re.match(r'Cycle\s+(\d+):', line, re.IGNORECASE)
            if cycle_match:
                current_cycle = int(cycle_match.group(1))
                logger.debug(f"Found Cycle {current_cycle}")
                if current_cycle not in points_by_cycle:
                    points_by_cycle[current_cycle] = []
                continue
            
            # Extract bullet points OR long lines (which are likely points without bullets)
            point_text = None
            
            # Check for various bullet formats
            if line.startswith('•'):
                point_text = re.sub(r'^•\s*', '', line).strip()
                logger.debug(f"Found bullet point (•) in Cycle {current_cycle}: '{point_text}'")
            
            elif line.startswith('-') and not line.startswith('--'):
                point_text = re.sub(r'^-\s*', '', line).strip()
                logger.debug(f"Found dash point (-) in Cycle {current_cycle}: '{point_text}'")
            
            elif line.startswith('*') and not line.startswith('**'):
                point_text = re.sub(r'^\*\s*', '', line).strip()
                logger.debug(f"Found asterisk point (*) in Cycle {current_cycle}: '{point_text}'")
            
            elif line.startswith('+'):
                point_text = re.sub(r'^\+\s*', '', line).strip()
                logger.debug(f"Found plus point (+) in Cycle {current_cycle}: '{point_text}'")
            
            elif re.match(r'^\d+\.', line):
                point_text = re.sub(r'^\d+\.\s*', '', line).strip()
                logger.debug(f"Found numbered point in Cycle {current_cycle}: '{point_text}'")
            
            # Fallback: If line is long (30+ chars) and meaningful, treat as point (only if in cycle)
            elif current_cycle is not None and len(line) >= 30 and not line.startswith(('=', '_', '-', 'Cycle')):
                point_text = line
                logger.debug(f"Found potential point (no bullet, long line) in Cycle {current_cycle}: '{line[:50]}...'")
            
            # Add point if found
            if point_text:
                all_points.append(point_text)
                if current_cycle is not None:
                    points_by_cycle[current_cycle].append(point_text)
        
        logger.debug(f"Total points found: {len(all_points)}")
        logger.debug(f"Points by cycle: {[(cycle, len(points)) for cycle, points in sorted(points_by_cycle.items())]}")
        
        return points_by_cycle, all_points
    
    def find_bookmark_paragraph(self, doc, bookmark_name):
        """Find the paragraph that contains a bookmark."""
        # Search all paragraphs for the bookmark
        for para in doc.paragraphs:
            # Check paragraph XML directly
            if bookmark_name in para._element.xml:
                return para
        
        # If not found in paragraphs, check all elements more thoroughly
        from docx.oxml import OxmlElement as OE
        for element in doc.element.iter():
            if bookmark_name in element.tag or (hasattr(element, 'attrib') and any(bookmark_name in str(v) for v in element.attrib.values())):
                # Find parent paragraph
                parent = element.getparent()
                while parent is not None:
                    if 'p' in parent.tag:  # paragraph element
                        # Find the paragraph object
                        for para in doc.paragraphs:
                            if para._element == parent:
                                return para
                    parent = parent.getparent()
        
        return None
    
    def copy_list_formatting(self, source_para, target_para):
        """
        Copy list/bullet formatting from source paragraph to target paragraph.
        This ensures bullets are properly applied. Includes fallback formatting.
        """
        try:
            source_pPr = source_para._element.get_or_add_pPr()
            target_pPr = target_para._element.get_or_add_pPr()
            
            # Look for numPr (numbering properties) element
            source_numPr = source_pPr.find(qn('w:numPr'))
            
            if source_numPr is not None:
                try:
                    # Deep copy the numbering properties
                    target_numPr = target_pPr.find(qn('w:numPr'))
                    if target_numPr is not None:
                        target_pPr.remove(target_numPr)
                    
                    # Copy the numPr element
                    new_numPr = deepcopy(source_numPr)
                    target_pPr.append(new_numPr)
                    logger.debug("Successfully copied list formatting (numPr) to new paragraph")
                except Exception as copy_error:
                    # Fallback: add manual bullet prefix if numPr copy fails
                    logger.debug(f"NumPr copy failed, applying fallback formatting: {str(copy_error)}")
                    self._apply_fallback_bullet_formatting(target_para)
            else:
                logger.debug("Source paragraph has no list formatting (numPr) to copy")
                
        except Exception as e:
            logger.warning(f"Could not copy list formatting, will apply fallback: {str(e)}")
            try:
                self._apply_fallback_bullet_formatting(target_para)
            except Exception as fallback_error:
                logger.debug(f"Fallback formatting also failed: {str(fallback_error)}")
    
    def _apply_fallback_bullet_formatting(self, para):
        """
        Apply fallback bullet formatting when XML manipulation fails.
        Prepends bullet character to paragraph text.
        """
        try:
            if para.runs:
                first_run = para.runs[0]
                if first_run.text and not first_run.text.startswith('•'):
                    first_run.text = '• ' + first_run.text
                    logger.debug("Applied fallback bullet prefix to paragraph")
        except Exception as e:
            logger.debug(f"Could not apply fallback bullet formatting: {str(e)}")
    
    def inject_points_into_resume(self, resume_bytes, processed_text, custom_mapping=None):
        """
        Inject extracted points into resume at bookmarks with flexible mapping.
        
        Args:
            resume_bytes: BytesIO object of the resume template
            processed_text: Processed text organized by cycles
            custom_mapping: Dict of {cycle_num: bookmark_name}. If None, auto-generates.
            
        Returns:
            Tuple of (BytesIO with updated resume, injection_details dict)
        """
        try:
            try:
                doc = Document(resume_bytes)
            except Exception as e:
                raise ValueError(f"Invalid or corrupted DOCX file: {str(e)}. Please check the resume template.")
            
            # Reset stream position for potential re-reads
            resume_bytes.seek(0)
            
            points_by_cycle, all_points = self.extract_points_by_heading(processed_text)
            
            if not all_points:
                raise ValueError("No points found in processed text. Check the format.")
            
            # Validate that we have actual points, not just empty cycles
            non_empty_cycles = {c: p for c, p in points_by_cycle.items() if p}
            if not non_empty_cycles:
                raise ValueError("No actual points found in any cycles. Cycles are empty.")
            
            # Auto-detect all bookmarks in the resume
            available_bookmarks = self.bookmark_manager.detect_bookmarks(resume_bytes)
            
            if not available_bookmarks:
                raise ValueError("No bookmarks found in resume template. Please add bookmarks first.")
            
            # Generate or use provided mapping
            if custom_mapping:
                cycle_to_bookmark = custom_mapping
                # Validate the custom mapping
                is_valid, error_msg = self.bookmark_manager.validate_mapping(
                    cycle_to_bookmark, available_bookmarks
                )
                if not is_valid:
                    raise ValueError(f"Invalid mapping: {error_msg}")
            else:
                # Auto-suggest mapping based on patterns and position
                num_cycles = len(points_by_cycle)
                cycle_to_bookmark = self.bookmark_manager.suggest_mappings(
                    available_bookmarks, num_cycles
                )
            
            if not cycle_to_bookmark:
                raise ValueError("Could not generate bookmark mappings. Please provide custom mapping.")
            
            logger.debug(f"Available bookmarks: {available_bookmarks}")
            logger.debug(f"Cycle to bookmark mapping: {cycle_to_bookmark}")
            logger.debug(f"Points by cycle: {points_by_cycle}")
            
            # Track injections for feedback
            injections = {}
            
            # Inject points for each cycle into its corresponding bookmark
            for cycle_num in sorted(points_by_cycle.keys()):
                if cycle_num not in cycle_to_bookmark:
                    logger.debug(f"Warning: Cycle {cycle_num} has no corresponding bookmark")
                    continue
                
                bookmark_name = cycle_to_bookmark[cycle_num]
                cycle_points = points_by_cycle[cycle_num]
                
                if not cycle_points:
                    logger.debug(f"Cycle {cycle_num} has no points to inject")
                    continue
                
                logger.debug(f"Injecting Cycle {cycle_num} points into {bookmark_name}")
                
                bookmark_para = self.find_bookmark_paragraph(doc, bookmark_name)
                
                if not bookmark_para:
                    # Bookmark not found, try to insert after the company section heading
                    company_section_name = bookmark_name.replace('_Responsibilities', '')
                    found = False
                    for para_idx, para in enumerate(doc.paragraphs):
                        if 'Responsibilities' in para.text and company_section_name in para.text:
                            bookmark_para = doc.paragraphs[para_idx + 1] if para_idx + 1 < len(doc.paragraphs) else para
                            found = True
                            logger.debug(f"Found {company_section_name} section at paragraph {para_idx}")
                            break
                    
                    if not found:
                        # Try any responsibility section
                        for para_idx, para in enumerate(doc.paragraphs):
                            if 'Responsibilities' in para.text:
                                bookmark_para = doc.paragraphs[min(para_idx + 2, len(doc.paragraphs) - 1)]
                                found = True
                                break
                
                if bookmark_para:
                    reference_para = bookmark_para
                    
                    # Get the parent element and insertion point
                    parent = reference_para._element.getparent()
                    ref_index = None
                    
                    # Try multiple strategies to find the reference index
                    try:
                        ref_index = list(parent).index(reference_para._element)
                    except ValueError:
                        pass
                    
                    # Fallback: find by matching element in paragraph list
                    if ref_index is None:
                        for idx, para in enumerate(doc.paragraphs):
                            if para._element == reference_para._element:
                                ref_index = idx
                                break
                    
                    # Last resort: use the end of the parent
                    if ref_index is None:
                        logger.warning("Could not find reference paragraph index, appending to end")
                        ref_index = max(0, len(list(parent)) - 1)
                    
                    # Add points for this cycle
                    for i, point_text in enumerate(cycle_points):
                        # Get the reference style - use the same style as reference
                        ref_style_name = reference_para.style.name if reference_para.style else 'Normal'
                        
                        # Add new paragraph directly to the document
                        new_para = doc.add_paragraph(point_text, style=ref_style_name)
                        
                        # Copy paragraph formatting from reference
                        ref_pformat = reference_para.paragraph_format
                        new_pformat = new_para.paragraph_format
                        
                        new_pformat.left_indent = ref_pformat.left_indent
                        new_pformat.first_line_indent = ref_pformat.first_line_indent
                        new_pformat.space_before = ref_pformat.space_before
                        new_pformat.space_after = ref_pformat.space_after
                        if ref_pformat.line_spacing:
                            new_pformat.line_spacing = ref_pformat.line_spacing
                        
                        # Copy font formatting from reference
                        if reference_para.runs and len(reference_para.runs) > 0:
                            ref_run = reference_para.runs[0]
                            for run in new_para.runs:
                                if ref_run.font.name:
                                    run.font.name = ref_run.font.name
                                if ref_run.font.size:
                                    run.font.size = ref_run.font.size
                                if ref_run.font.bold:
                                    run.font.bold = ref_run.font.bold
                                if ref_run.font.italic:
                                    run.font.italic = ref_run.font.italic
                                if ref_run.font.color.rgb:
                                    run.font.color.rgb = ref_run.font.color.rgb
                        
                        # CRITICAL: Copy list/bullet formatting from reference
                        self.copy_list_formatting(reference_para, new_para)
                        
                        # Now move the paragraph to the correct position (after reference)
                        new_parent = new_para._element.getparent()
                        
                        # Remove from the end where it was added
                        new_parent.remove(new_para._element)
                        
                        # Insert at the correct position in the parent
                        try:
                            ref_index = list(parent).index(reference_para._element)
                            parent.insert(ref_index + 1 + i, new_para._element)
                        except (ValueError, IndexError) as e:
                            logger.warning(f"Could not insert at position, appending instead: {str(e)}")
                            parent.append(new_para._element)
                    
                    injections[bookmark_name] = len(cycle_points)
                    logger.debug(f"Successfully injected {len(cycle_points)} points into {bookmark_name}")
            
            if not injections:
                raise ValueError("Failed to inject points. No valid insertion points found.")
            
            # Save to BytesIO
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return output, injections
            
        except Exception as e:
            raise Exception(f"Error injecting points into resume: {str(e)}")
    
    def get_available_bookmarks(self, resume_bytes):
        """Get list of available bookmarks in resume template (delegated to BookmarkManager)."""
        return self.bookmark_manager.detect_bookmarks(resume_bytes)
