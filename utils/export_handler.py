import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT

class ExportHandler:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Setup custom paragraph styles for PDF generation."""
        self.styles.add(ParagraphStyle(
            name='CycleHeading',
            parent=self.styles['Heading1'],
            fontSize=14,
            textColor='#000000',
            spaceAfter=12,
            spaceBefore=6
        ))
        self.styles.add(ParagraphStyle(
            name='ContentText',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leading=14
        ))

    def generate_docx(self, content):
        """Generate a DOCX file from the processed text."""
        document = Document()

        # Split the content by newlines
        lines = content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Add headings and regular content
            if line.startswith('Cycle'):
                document.add_heading(line, level=1)
            else:
                document.add_paragraph(line)

        # Save the document to a BytesIO object
        docx_file = io.BytesIO()
        document.save(docx_file)
        docx_file.seek(0)

        return docx_file

    def generate_pdf(self, content):
        """Generate a PDF file from the processed text with proper text wrapping and pagination."""
        pdf_file = io.BytesIO()
        doc = SimpleDocTemplate(
            pdf_file,
            pagesize=letter,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=50
        )

        # Build story of elements
        story = []
        lines = content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1 * inch))
                continue

            # Determine line type and formatting
            if line.startswith('Cycle'):
                story.append(Paragraph(line, self.styles['CycleHeading']))
            else:
                story.append(Paragraph(line, self.styles['ContentText']))

        # Build PDF
        doc.build(story)
        pdf_file.seek(0)
        return pdf_file