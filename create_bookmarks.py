from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark_to_paragraph(paragraph, bookmark_name):
    """Add a bookmark to a paragraph."""
    run = paragraph.add_run()
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark_name)
    run._element.addprevious(start)
    
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    run._element.addnext(end)

def add_bookmarks_to_resume(resume_path, output_path):
    """Add bookmarks to all company responsibility sections."""
    doc = Document(resume_path)
    
    # Company section definitions with their responsibility ranges
    companies = [
        {
            'name': 'KPMG',
            'bookmark': 'KPMG_Responsibilities',
            'last_para': 54
        },
        {
            'name': 'CVS',
            'bookmark': 'CVS_Responsibilities',
            'last_para': 91
        },
        {
            'name': 'Harland',
            'bookmark': 'Harland_Responsibilities',
            'last_para': 105
        },
        {
            'name': 'First Citizen Bank',
            'bookmark': 'FirstCitizensBank_Responsibilities',
            'last_para': 114
        }
    ]
    
    # Add bookmarks at the end of each responsibility section
    for company in companies:
        para = doc.paragraphs[company['last_para']]
        add_bookmark_to_paragraph(para, company['bookmark'])
        print(f"Added bookmark '{company['bookmark']}' at Para {company['last_para']}")
    
    # Save the document
    doc.save(output_path)
    print(f"\nResume saved with bookmarks to: {output_path}")

if __name__ == "__main__":
    input_path = r'C:\Users\12shi\Downloads\Lead Engineer.docx'
    output_path = r'C:\Users\12shi\Downloads\Lead Engineer_with_bookmarks.docx'
    
    add_bookmarks_to_resume(input_path, output_path)
    print("\nBookmarks added successfully!")
