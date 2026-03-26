from docx import Document

doc = Document(r'C:\Users\12shi\Downloads\Lead Engineer.docx')

print('=== DOCUMENT STRUCTURE ===\n')
print(f'Total paragraphs: {len(doc.paragraphs)}\n')

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        text_preview = para.text[:100] if len(para.text) > 100 else para.text
        print(f'Para {i}: {text_preview}')
        print(f'  Style: {para.style.name}')
        if para.runs and para.runs[0].font.name:
            print(f'  Font: {para.runs[0].font.name}')
        print()
