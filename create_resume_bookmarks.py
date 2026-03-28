#!/usr/bin/env python3
"""
Bookmark Creator for Resume Templates
Helps users add bookmarks to their resume for the Resume Injection feature.
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import sys
from pathlib import Path

def add_bookmark_to_paragraph(paragraph, bookmark_name):
    """Add a bookmark to a paragraph."""
    run = paragraph.add_run()
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark_name)
    run._element.addprevious(start)
    
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    run._element.addnext(end)

def interactive_bookmark_creation():
    """Guide user through creating bookmarks in their resume."""
    
    print("\n" + "="*60)
    print("RESUME BOOKMARK CREATOR")
    print("="*60)
    print("\nThis tool helps you add bookmarks to your resume for")
    print("the Resume Injection feature.\n")
    
    # Get resume file path
    while True:
        resume_path = input("Enter path to your resume (DOCX file): ").strip()
        
        if not resume_path:
            print("❌ Path cannot be empty")
            continue
        
        try:
            if not Path(resume_path).exists():
                print(f"❌ File not found: {resume_path}")
                continue
            
            doc = Document(resume_path)
            print(f"✅ Successfully loaded: {resume_path}\n")
            break
        except Exception as e:
            print(f"❌ Error loading file: {e}")
            continue
    
    # Find sections
    print("Finding 'Responsibilities:' sections in your resume...\n")
    responsibility_sections = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "Responsibilities:":
            # Find the company name by looking backwards
            company_name = "Unknown"
            for j in range(i-1, max(0, i-5), -1):
                prev_text = doc.paragraphs[j].text.strip()
                if prev_text and 'Client' in prev_text:
                    company_name = prev_text
                    break
            
            responsibility_sections.append({
                'company': company_name,
                'responsibility_header_para': i,
                'para_object': para
            })
    
    if not responsibility_sections:
        print("❌ No 'Responsibilities:' sections found.")
        print("Ensure your resume has this structure:")
        print("  Company Name")
        print("  Position Title")
        print("  Responsibilities:")
        print("  • Point 1")
        print("  • Point 2")
        return
    
    print(f"Found {len(responsibility_sections)} Responsibilities section(s):\n")
    
    # Show options and let user choose
    for idx, section in enumerate(responsibility_sections, 1):
        print(f"{idx}. {section['company']}")
    
    # Create bookmarks
    bookmarks_to_add = []
    
    for idx, section in enumerate(responsibility_sections, 1):
        company = section['company']
        
        # Generate suggested bookmark name
        suggested_name = company.replace('Client - ', '').replace(' – ', '_').replace(' - ', '_').replace(' ', '')
        suggested_name = f"{suggested_name}_Responsibilities"
        
        print(f"\n{idx}. {company}")
        custom_name = input(f"   Bookmark name [{suggested_name}]: ").strip()
        
        bookmark_name = custom_name if custom_name else suggested_name
        bookmarks_to_add.append({
            'name': bookmark_name,
            'company': company,
            'para': section['para_object']
        })
    
    # Confirm before adding
    print("\n" + "="*60)
    print("BOOKMARK PLACEMENT OPTIONS")
    print("="*60)
    print("""
1. 🏷️  AT HEADER (Responsibilities:)
   - Injects right after "Responsibilities:" line
   - New points appear first in the section

2. 📍 AT END OF SECTION
   - Injects after all existing bullet points
   - New points appear at the end (recommended)
   ✨ RECOMMENDED for appending new points to existing ones

3. ❌ CANCEL """)
    
    placement = input("\nChoose placement (1, 2, or 3): ").strip()
    
    if placement == '3':
        print("❌ Cancelled. No changes made.")
        return
    elif placement not in ['1', '2']:
        print("❌ Invalid choice")
        return
    
    print("\n" + "="*60)
    print("SUMMARY - Bookmarks to Add:")
    print("="*60)
    
    placement_type = "AT HEADER" if placement == '1' else "AT END OF SECTION"
    print(f"Placement: {placement_type}")
    print()
    
    for bookmark in bookmarks_to_add:
        print(f"✓ '{bookmark['name']}' at {bookmark['company']}")
    
    confirm = input("\nProceed with adding these bookmarks? (yes/no): ").strip().lower()
    
    if confirm not in ['yes', 'y']:
        print("❌ Cancelled. No changes made.")
        return
    
    # Add bookmarks
    for idx, section in enumerate(responsibility_sections):
        bookmark = bookmarks_to_add[idx]
        
        try:
            if placement == '1':
                # Add at header
                add_bookmark_to_paragraph(section['para_object'], bookmark['name'])
                print(f"✅ Added bookmark at header: {bookmark['name']}")
            else:
                # Find last bullet point in this section and add after it
                responsibility_idx = section['responsibility_header_para']
                
                # Find the last non-empty paragraph in this section
                last_point_idx = responsibility_idx
                for j in range(responsibility_idx + 1, len(doc.paragraphs)):
                    next_para = doc.paragraphs[j]
                    text = next_para.text.strip()
                    
                    # Stop if we hit a new section (company name, "Accountabilities:", etc)
                    if (text.startswith('Sr.') or text.startswith('Lead') or 
                        text.startswith('Senior') or 'Client' in text or
                        text.endswith('LLC') or text.endswith('Inc') or
                        text.endswith('Group') or text == 'Accountabilities:'):
                        break
                    
                    if text:  # Non-empty paragraph
                        last_point_idx = j
                
                # Add bookmark at end of section
                add_bookmark_to_paragraph(doc.paragraphs[last_point_idx], bookmark['name'])
                print(f"✅ Added bookmark at end of section: {bookmark['name']}")
        
        except Exception as e:
            print(f"❌ Error adding bookmark '{bookmark['name']}': {e}")
    
    # Save the document
    output_path = resume_path.replace('.docx', '_with_bookmarks.docx')
    
    try:
        doc.save(output_path)
        print(f"\n✅ Successfully saved resume with bookmarks!")
        print(f"   Output: {output_path}")
        print(f"\n💡 Placement details:")
        if placement == '1':
            print("   Bookmarks placed at 'Responsibilities:' headers")
            print("   New points will appear immediately after headers")
        else:
            print("   Bookmarks placed at end of each section")
            print("   New points will appear after all existing points (recommended)")
        print(f"\n💡 Tip: Use this file as your template for Resume Injection.")
    except Exception as e:
        print(f"❌ Error saving file: {e}")

if __name__ == "__main__":
    try:
        interactive_bookmark_creation()
    except KeyboardInterrupt:
        print("\n\n❌ Cancelled by user")
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ Unexpected error: {e}")
        sys.exit(1)
