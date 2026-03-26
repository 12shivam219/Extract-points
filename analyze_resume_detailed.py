from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph."""
    run = paragraph.add_run()
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), '0')
    start.set(qn('w:name'), bookmark_name)
    run._element.addprevious(start)
    
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), '0')
    run._element.addnext(end)

def analyze_resume():
    """Analyze resume and find all company sections."""
    doc = Document(r'C:\Users\12shi\Downloads\Lead Engineer.docx')
    
    print("=== COMPANY SECTIONS FOUND ===\n")
    
    companies = []
    current_company = None
    responsibility_start = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Look for company lines (contain "Client -" or company patterns)
        if 'Client -' in text and '–' in text:
            if current_company:
                companies.append({
                    'name': current_company['name'],
                    'company_para': current_company['company_para'],
                    'resp_start': current_company['resp_start'],
                    'resp_end': i - 2  # Last para before new company
                })
            
            current_company = {
                'name': text,
                'company_para': i,
                'resp_start': None
            }
            print(f"Company: {text} (Para {i})")
        
        # Look for "Responsibilities:" line
        if text == "Responsibilities:" and current_company:
            current_company['resp_start'] = i + 1
            print(f"  Responsibilities start: Para {i + 1}")
    
    # Add last company
    if current_company:
        companies.append({
            'name': current_company['name'],
            'company_para': current_company['company_para'],
            'resp_start': current_company['resp_start'],
            'resp_end': len(doc.paragraphs) - 1
        })
    
    print("\n=== SUMMARY ===")
    for company in companies:
        print(f"\n{company['name']}")
        print(f"  Responsibilities: Para {company['resp_start']} to {company['resp_end']}")
    
    return doc, companies

if __name__ == "__main__":
    doc, companies = analyze_resume()
